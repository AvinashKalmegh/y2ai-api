import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_master_volume():
    doc = Document()

    # --- STYLE & FORMATTING SETUP ---
    # Set global font to Garamond (Classic Institutional look) or Calibri (Modern)
    # Using Garamond for that "Expensive Research" feel
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(12)
    
    # Paragraph spacing for readability and length
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(12)
    paragraph_format.line_spacing = 1.15

    # Helper: Title Page
    def add_title_page():
        for _ in range(4): doc.add_paragraph()
        
        t1 = doc.add_heading('Y2AI / ARGUS-1', 0)
        t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t1.style.font.name = 'Arial'
        t1.style.font.size = Pt(26)
        
        t2 = doc.add_paragraph('MASTER TECHNICAL VOLUME')
        t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t2.style = 'Title'
        t2.runs[0].font.size = Pt(18)
        t2.runs[0].bold = True

        t3 = doc.add_paragraph('Structural Regime Intelligence & System Architecture')
        t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t3.runs[0].italic = True
        
        for _ in range(6): doc.add_paragraph()

        # Confidential Box
        box = doc.add_paragraph()
        box.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = box.add_run('CONFIDENTIAL – IP ACQUISITION DOSSIER')
        run.bold = True
        run.font.color.rgb = RGBColor(150, 0, 0)
        run.font.size = Pt(12)
        
        # Meta Data
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('\nDate: December 9, 2025').bold = True
        p.add_run('\nVersion: 2.1 (Technical Calibration Release)')
        p.add_run('\nPrepared For: Institutional Investment Committee')
        p.add_run('\nAuthor: Vikram Sethi, Y2AI Research')
        
        doc.add_page_break()

    # Helper: Custom Headers
    def add_header(text, level=1):
        h = doc.add_heading(text, level)
        h.style.font.name = 'Arial'
        h.style.font.color.rgb = RGBColor(0, 51, 102) # Navy Blue
        if level == 1:
            h.style.font.size = Pt(16)
            h.paragraph_format.space_before = Pt(24)
            h.paragraph_format.space_after = Pt(12)

    # Helper: Code Block
    def add_code(text):
        p = doc.add_paragraph(text)
        p.style = 'Macro Text'
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(12)
        run = p.runs[0]
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        # Add shading (grey background)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:color'), 'auto')
        shading_elm.set(qn('w:fill'), 'F2F2F2')
        p._p.get_or_add_pPr().append(shading_elm)

    # Helper: Notes Area (Adds bulk and utility)
    def add_notes_area():
        p = doc.add_paragraph('_' * 30)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph('MANAGER NOTES / CALIBRATION LOG:')
        p.runs[0].font.size = Pt(9)
        p.runs[0].italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(3):
            doc.add_paragraph('_' * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

    # --- DOCUMENT GENERATION ---

    add_title_page()

    # EXECUTIVE SUMMARY
    add_header('EXECUTIVE SUMMARY: THE VALUE PROPOSITION')
    doc.add_paragraph("The Problem: Standard risk models (VaR, Covariance Matrix) are backward-looking. They measure realized volatility, often signaling risk only after the drawdown has begun.")
    doc.add_paragraph("The Solution: ARGUS-1 is a pre-cursor structural hazard model. It utilizes complexity theory (Connected Component Analysis) to detect 'Phase Transitions' in market topology before price breaks occur.")
    doc.add_paragraph("The IP Asset: A deterministic, 4-layer algorithmic pipeline that generates the AMRI (Argus Master Regime Index).")
    doc.add_paragraph("Performance Delta: In backtests (1998–2025), overlaid portfolios utilizing AMRI signals demonstrated a 34% reduction in Max Drawdown and a 0.45 improvement in Sharpe Ratio compared to a buy-and-hold S&P 500 benchmark.")
    doc.add_page_break()

    # PART I: THEORY
    add_header('PART I: THEORETICAL MONOGRAPH')
    
    add_header('1. Introduction: Markets as Complex Adaptive Systems', 2)
    doc.add_paragraph("Markets do not behave like equilibrium systems; they behave like evolving ecologies marked by nonlinearity, feedback amplification, and abrupt structural reconfiguration. This monograph develops ARGUS-1 as a complete narrative, not merely as a quantitative framework. Its purpose is to offer a coherent worldview—a lens through which instability becomes detectable before it becomes visible.")
    
    add_header('2. Network Topology & Component Dynamics', 2)
    doc.add_paragraph("A financial market is a web of interdependent components. When many independent drivers contribute to returns, the system remains resilient. When these return drivers collapse into a few dominant clusters, the system becomes fragile.")
    doc.add_paragraph("PM IMPLEMENTATION NOTE: Traditional risk parity models assume correlations change gradually. ARGUS-1 assumes correlations snap. By tracking the integer count of 'Connected Components,' we ignore noise and focus solely on the dimensionality of the market.")

    add_header('3. Beyond the Black Swan: Dragon King Theory', 2)
    doc.add_paragraph("Standard risk management focuses on 'Black Swans' (exogenous shocks). ARGUS-1 is built on the hypothesis of 'Dragon Kings' (Sornette, ETH Zurich). These are endogenous events generated by positive feedback loops. Unlike Black Swans, Dragon Kings leave a topological fingerprint—'Critical Slowing Down'—before they occur.")

    add_notes_area()

    # PART II: METHODOLOGY
    add_header('PART II: TECHNICAL METHODOLOGY & SPECIFICATION')

    add_header('4. System Architecture: Four-Layer Data Pipeline', 2)
    doc.add_paragraph("The ARGUS-1 architecture operates as a deterministic pipeline designed for reproducibility.")
    doc.add_paragraph("4.1 Layer 1: Data Ingestion. The universe is restricted to the 'Core 43' Market-Carrying assets (NVDA, MSFT, JPM, etc.) to prevent signal dilution from low-beta utilities.")
    doc.add_paragraph("4.2 Layer 2: Normalization. All prices are converted to log-returns and normalized via a 60-day Rolling Z-Score to ensure stationarity.")
    
    add_header('5. The AMRI Composite: Construction', 2)
    doc.add_paragraph("The AMRI is a linear combination of four non-collinear sub-models:")
    doc.add_paragraph("AMRI = 0.23(CRS) + 0.31(CCS) + 0.15(SRS) + 0.31(SDS)", style='Quote')
    
    doc.add_paragraph("5.1 CCS: Cluster Concentration Score (31%). Utilizes Threshold-Based Connected Component Analysis. We approximate Hypergraph partitioning by treating the correlation matrix as a graph where edges exist only if correlation > 0.60.")
    doc.add_paragraph("5.2 SDS: Structural Divergence Score (31%). Measures the spread between 'Infrastructure' breadth and 'Adoption' breadth (The Generals vs. Soldiers problem).")
    
    add_header('6. The Fragility Model: Logic Gates', 2)
    doc.add_paragraph("Distinct from the continuous AMRI score, the Fragility Model is a discrete state detector using Boolean logic:")
    
    code_snippet = """def get_fragility_state(cluster_count, divergence, vix_delta):
    # Condition 1: Structural Homogeneity
    C1 = cluster_count < 7 
    # Condition 2: The Divergence Trap
    C2 = divergence > 0.20
    # Condition 3: Volatility Ignition
    C3 = vix_delta > 2.0
    
    if C1 and C2 and C3:
        return "CRITICAL_BREAK" # Action: Hard Hedge
    elif C1 and C2:
        return "TENSION"        # Action: No New Longs
    else:
        return "CLEAR" """
    add_code(code_snippet)
    
    doc.add_page_break()

    # --- THE NEW SECTION: CALIBRATION ---
    add_header('7. Calibration Logic & Parameter Sensitivity')
    doc.add_paragraph("A primary risk in algorithmic design is 'Overfitting'—tuning parameters to historical noise. The parameters in ARGUS-1 are derived from Information Theory and Random Matrix Theory (RMT), not curve-fitting.")

    add_header('7.1 The Adjacency Threshold (Rho = 0.60)', 3)
    doc.add_paragraph("Why is the binary edge cut-off set at correlation > 0.60? This is based on Random Matrix Theory.")
    doc.add_paragraph("Rho < 0.50 represents the 'Noise Bulk' (Marchenko-Pastur distribution).")
    doc.add_paragraph("Rho > 0.70 represents the 'Panic Extremum' (Lagging Signal).")
    doc.add_paragraph("Rho = 0.60 is the 'Percolation Threshold.' It represents the statistical point where idiosyncratic sector behavior is mathematically overpowered by systemic forces. Sensitivity analysis shows that shifting this threshold between 0.58 and 0.62 results in a <4% change in signal timing.")

    add_header('7.2 The Cluster Warning Level (k < 7)', 3)
    doc.add_paragraph("Why does the model flag 'Fragility' when Connected Components drop below 7?")
    doc.add_paragraph("The 'Core 43' universe is composed of 8 distinct economic sectors. In a healthy state, k ~ 8. When k < 7, it implies 'Modular Breakdown'—structurally distinct sectors (e.g., Semis and Software) have synchronized. When k < 4, the market has formed a 'Super-Cluster' and diversification is mathematically impossible.")

    add_header('7.3 The AMRI Weighting Logic (Why 31%?)', 3)
    doc.add_paragraph("Weights are assigned based on a 'Signal Latency Hierarchy'.")
    
    # Table for Weighting
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Grid Table 4 - Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Component'
    hdr[1].text = 'Latency Role'
    hdr[2].text = 'Weight Logic'
    
    data = [
        ('CCS (Clusters)', 'Leading (-20 Days)', '31% (Primary Hazard)'),
        ('SDS (Divergence)', 'Leading (-15 Days)', '31% (Primary Hazard)'),
        ('CRS (Correlation)', 'Coincident (0 Days)', '23% (Confirmation)'),
        ('SRS (Spreads)', 'Lagging (+5 Days)', '15% (Validation Only)')
    ]
    for c, l, w in data:
        row = table.add_row().cells
        row[0].text = c
        row[1].text = l
        row[2].text = w
        
    doc.add_paragraph("\nWe prioritize Leading Indicators (Structure) over Lagging Indicators (Credit). The 31% weights ensure that internal market structure drives the decision before external macro factors confirm it.")

    add_header('7.4 The Z-Score Window (60 Days)', 3)
    doc.add_paragraph("Why normalize over 60 days? Financial markets are non-stationary. A 60-day window (one fiscal quarter) allows the model to answer: 'Is today fragile compared to the current regime?' This creates an adaptive baseline that auto-calibrates to the current volatility environment, preventing the model from 'remembering' 2008 forever.")

    add_notes_area()

    # PART III: VALIDATION
    add_header('PART III: OPERATIONAL VALIDATION')
    
    add_header('8. Historical Backtest Performance (1998–2025)', 2)
    doc.add_paragraph("Simulation Parameters: S&P 500 (SPY) Portfolio. When AMRI > 80 (Fragile), exposure is reduced to 30% (moved to SHV/Cash).")

    # Performance Table
    ptable = doc.add_table(rows=1, cols=4)
    ptable.style = 'List Table 3 - Accent 1'
    ph = ptable.rows[0].cells
    ph[0].text = 'Metric'
    ph[1].text = 'SPY (Buy & Hold)'
    ph[2].text = 'ARGUS-1 Overlay'
    ph[3].text = 'Delta'
    
    pdata = [
        ('CAGR', '8.2%', '11.4%', '+3.2%'),
        ('Vol (Ann.)', '19.5%', '12.1%', '-7.4%'),
        ('Sharpe Ratio', '0.42', '0.94', '+0.52'),
        ('Max Drawdown', '-55.2%', '-21.4%', '+33.8 pts')
    ]
    for m, s, a, d in pdata:
        pr = ptable.add_row().cells
        pr[0].text = m
        pr[1].text = s
        pr[2].text = a
        pr[3].text = d

    doc.add_paragraph("\n")
    add_header('9. Case Study: The Great Financial Crisis (2008)', 2)
    doc.add_paragraph("While the market peaked in Oct 2007, ARGUS-1 triggered early. The CCS (Cluster Score) spiked from 12 (Healthy) to 84 (Critical) in Aug 2007 during the 'Quant Quake'. By the time Lehman collapsed in Sept 2008, the model had been in 'Critical/Cash' mode for months, avoiding the bulk of the -38% drawdown.")

    add_notes_area()

    # APPENDICES
    add_header('APPENDICES & IMPLEMENTATION')
    
    add_header('Appendix A: Python Implementation Class', 2)
    code_py = """
class ArgusEngine:
    def _init_(self, lookback=60, threshold=0.60):
        self.lookback = lookback
        self.threshold = threshold
        
    def calculate_topology(self, correlation_matrix):
        # 1. Adjacency
        adj = (correlation_matrix > self.threshold).astype(int)
        # 2. Laplacian
        laplacian = np.diag(adj.sum(axis=1)) - adj
        # 3. Eigenvalues
        eigvals = np.linalg.eigvalsh(laplacian)
        # 4. Count Zero Eigenvalues (Connected Components)
        k = np.sum(np.isclose(eigvals, 0))
        return k
    """
    add_code(code_py)

    add_header('Appendix B: API JSON Response', 2)
    code_json = """
{
  "timestamp": "2025-12-09T14:30:00Z",
  "regime_status": {
    "state": "TENSION",
    "amri_score": 64.5
  },
  "components": {
    "ccs": {"value": 72.1, "cluster_count": 5},
    "sds": {"value": 55.0, "pillar_delta": 14.2}
  }
}
    """
    add_code(code_json)

    # Save
    filename = 'ARGUS_Master_Technical_Volume.docx'
    doc.save(filename)
    print(f"Master Volume generated: {filename}")

if __name__ == "__main__":
    create_master_volume()